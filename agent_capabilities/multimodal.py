"""
Helix Agent Multi-Modal Capabilities
====================================
Image processing, document handling, code analysis, and data visualization
capabilities for Helix agents.

Features:
- Image understanding and analysis
- Document processing (PDF, DOCX, etc.)
- Code analysis and generation
- Data visualization
- OCR and text extraction
"""

import base64
import io
import json
import os
import re
from dataclasses import dataclass, field
from datetime import UTC, datetime
from enum import Enum
from pathlib import Path
from typing import Any

# Optional imports with fallbacks
try:
    from PIL import Image
    from PIL.ExifTags import TAGS as EXIF_TAGS

    HAS_PIL = True
except ImportError:
    Image = None  # type: ignore
    EXIF_TAGS = {}
    HAS_PIL = False

try:
    import pytesseract

    HAS_TESSERACT = True
except ImportError:
    pytesseract = None  # type: ignore
    HAS_TESSERACT = False

try:
    import fitz  # PyMuPDF

    HAS_PYMUPDF = True
except ImportError:
    fitz = None  # type: ignore
    HAS_PYMUPDF = False

try:
    from docx import Document as DocxDocument

    HAS_DOCX = True
except ImportError:
    DocxDocument = None  # type: ignore
    HAS_DOCX = False

try:
    import pandas as pd

    HAS_PANDAS = True
except ImportError:
    pd = None  # type: ignore
    HAS_PANDAS = False

try:
    import numpy as np

    HAS_NUMPY = True
except ImportError:
    np = None  # type: ignore
    HAS_NUMPY = False

try:
    import matplotlib

    matplotlib.use("Agg")  # Non-interactive backend
    import matplotlib.pyplot as plt

    HAS_MATPLOTLIB = True
except ImportError:
    HAS_MATPLOTLIB = False

try:
    import ast

    HAS_AST = True
except ImportError:
    ast = None  # type: ignore
    HAS_AST = False

# Logging
import logging

logger = logging.getLogger(__name__)


# ============================================================================
# FEATURE FLAGS - ML/Audio Capabilities
# ============================================================================
# Controls whether API-dependent features raise errors or return graceful
# fallbacks when API keys are missing.

# Audio transcription via OpenAI Whisper API (requires OPENAI_API_KEY)
# Defaults to "auto" — enabled when OPENAI_API_KEY is present.
_audio_flag = os.getenv("FEATURE_AUDIO_TRANSCRIPTION", "auto").lower()
FEATURE_AUDIO_TRANSCRIPTION = _audio_flag == "true" or (_audio_flag == "auto" and bool(os.getenv("OPENAI_API_KEY")))

# Image generation via DALL-E 3 / Grok (requires OPENAI_API_KEY or XAI_API_KEY)
# Defaults to "auto" — enabled when any image-capable API key is present.
_image_flag = os.getenv("FEATURE_IMAGE_GENERATION", "auto").lower()
FEATURE_IMAGE_GENERATION = _image_flag == "true" or (
    _image_flag == "auto" and bool(os.getenv("OPENAI_API_KEY") or os.getenv("XAI_API_KEY"))
)

# Video generation via Grok (requires XAI_API_KEY)
_video_flag = os.getenv("FEATURE_VIDEO_GENERATION", "auto").lower()
FEATURE_VIDEO_GENERATION = _video_flag == "true" or (_video_flag == "auto" and bool(os.getenv("XAI_API_KEY")))

# When a feature flag resolves to False and the API key is missing, the method
# returns a graceful fallback instead of raising. Set the flag to "true" to get
# explicit errors when the key is not configured.


class ImageFormat(Enum):
    """Supported image formats"""

    PNG = "png"
    JPEG = "jpeg"
    GIF = "gif"
    WEBP = "webp"
    BMP = "bmp"


class DocumentFormat(Enum):
    """Supported document formats"""

    PDF = "pdf"
    DOCX = "docx"
    TXT = "txt"
    MD = "md"
    HTML = "html"
    CSV = "csv"
    JSON = "json"
    XML = "xml"


@dataclass
class ImageAnalysisResult:
    """Result of image analysis"""

    width: int
    height: int
    format: str
    mode: str
    file_size: int
    has_transparency: bool
    dominant_colors: list[str] = field(default_factory=list)
    text_content: str | None = None
    objects_detected: list[dict[str, Any]] = field(default_factory=list)
    metadata: dict[str, Any] = field(default_factory=dict)

    def to_dict(self) -> dict[str, Any]:
        return {
            "width": self.width,
            "height": self.height,
            "format": self.format,
            "mode": self.mode,
            "file_size": self.file_size,
            "has_transparency": self.has_transparency,
            "dominant_colors": self.dominant_colors,
            "text_content": self.text_content,
            "objects_detected": self.objects_detected,
            "metadata": self.metadata,
        }


@dataclass
class DocumentContent:
    """Extracted document content"""

    text: str
    pages: int
    word_count: int
    format: str
    metadata: dict[str, Any] = field(default_factory=dict)
    sections: list[dict[str, Any]] = field(default_factory=list)
    tables: list[list[list[str]]] = field(default_factory=list)
    images: list[str] = field(default_factory=list)

    def to_dict(self) -> dict[str, Any]:
        return {
            "text": self.text,
            "pages": self.pages,
            "word_count": self.word_count,
            "format": self.format,
            "metadata": self.metadata,
            "sections": self.sections,
            "tables": self.tables,
            "image_count": len(self.images),
        }


@dataclass
class CodeAnalysisResult:
    """Result of code analysis"""

    language: str
    lines_of_code: int
    functions: list[dict[str, Any]]
    classes: list[dict[str, Any]]
    imports: list[str]
    complexity_score: float
    issues: list[dict[str, Any]]
    suggestions: list[str]
    documentation_coverage: float

    def to_dict(self) -> dict[str, Any]:
        return {
            "language": self.language,
            "lines_of_code": self.lines_of_code,
            "functions": self.functions,
            "classes": self.classes,
            "imports": self.imports,
            "complexity_score": self.complexity_score,
            "issues": self.issues,
            "suggestions": self.suggestions,
            "documentation_coverage": self.documentation_coverage,
        }


class ImageProcessor:
    """Process and analyze images"""

    def __init__(self):
        self.supported_formats = list(ImageFormat)

    async def analyze(
        self,
        image_data: bytes | str | Path,
        extract_text: bool = True,
        detect_objects: bool = False,
    ) -> ImageAnalysisResult:
        """Analyze an image"""
        if not HAS_PIL:
            return ImageAnalysisResult(
                width=0,
                height=0,
                format="unknown",
                mode="unknown",
                file_size=0,
                has_transparency=False,
                metadata={"error": "PIL not installed"},
            )

        # Load image
        if isinstance(image_data, (str, Path)):
            if str(image_data).startswith(("http://", "https://")):
                # Download from URL
                import aiohttp

                async with aiohttp.ClientSession() as session, session.get(str(image_data)) as response:
                    image_bytes = await response.read()
                img = Image.open(io.BytesIO(image_bytes))
                file_size = len(image_bytes)
            else:
                img = Image.open(image_data)
                file_size = Path(image_data).stat().st_size
        elif isinstance(image_data, bytes):
            img = Image.open(io.BytesIO(image_data))
            file_size = len(image_data)
        else:
            raise ValueError("Invalid image data type")

        # Basic analysis
        result = ImageAnalysisResult(
            width=img.width,
            height=img.height,
            format=img.format or "unknown",
            mode=img.mode,
            file_size=file_size,
            has_transparency=img.mode in ("RGBA", "LA", "P"),
        )

        # Extract dominant colors
        try:
            # Use PIL to extract dominant colors
            quantized = img.quantize(colors=8)  # Get 8 dominant colors
            palette = quantized.getpalette()
            if palette:
                # Convert RGB tuples to hex strings
                result.dominant_colors = [
                    f"#{palette[i]:02x}{palette[i+1]:02x}{palette[i+2]:02x}"
                    for i in range(0, min(24, len(palette)), 3)  # First 8 colors
                ]
        except Exception as e:
            logger.debug("Color extraction failed: %s", e)
            result.dominant_colors = []

        # OCR text extraction
        if extract_text and HAS_TESSERACT:
            try:
                result.text_content = pytesseract.image_to_string(img)
            except Exception as e:
                logger.debug("OCR text extraction failed: %s", e)
                result.text_content = None

        # Get EXIF metadata
        try:
            exif = img.getexif() if hasattr(img, "getexif") else None
            if exif:
                result.metadata["exif"] = {k: str(v) for k, v in exif.items() if isinstance(v, (str, int, float))}
        except Exception as e:
            logger.debug("Failed to extract EXIF metadata: %s", e)

        return result

    def _get_dominant_colors(self, img: "Image.Image", num_colors: int = 5) -> list[str]:
        """Extract dominant colors from image"""
        # Resize for faster processing
        img_small = img.copy()
        img_small.thumbnail((100, 100))

        # Convert to RGB if necessary
        if img_small.mode != "RGB":
            img_small = img_small.convert("RGB")

        # Get colors
        colors = img_small.getcolors(maxcolors=10000)
        if not colors:
            return []

        # Sort by frequency
        colors.sort(key=lambda x: x[0], reverse=True)

        # Convert to hex
        hex_colors = []
        for count, rgb in colors[:num_colors]:
            hex_color = "#{:02x}{:02x}{:02x}".format(*rgb)
            hex_colors.append(hex_color)

        return hex_colors

    async def resize(
        self,
        image_data: bytes | Path,
        width: int | None = None,
        height: int | None = None,
        maintain_aspect: bool = True,
    ) -> bytes:
        """Resize an image"""
        if not HAS_PIL:
            raise RuntimeError("PIL not installed")

        if isinstance(image_data, Path):
            img = Image.open(image_data)
        else:
            img = Image.open(io.BytesIO(image_data))

        if maintain_aspect:
            if width and height:
                img.thumbnail((width, height))
            elif width:
                ratio = width / img.width
                img = img.resize((width, int(img.height * ratio)))
            elif height:
                ratio = height / img.height
                img = img.resize((int(img.width * ratio), height))
        else:
            if width and height:
                img = img.resize((width, height))

        output = io.BytesIO()
        img.save(output, format=img.format or "PNG")
        return output.getvalue()

    async def convert(self, image_data: bytes | Path, target_format: ImageFormat) -> bytes:
        """Convert image to different format"""
        if not HAS_PIL:
            raise RuntimeError("PIL not installed")

        if isinstance(image_data, Path):
            img = Image.open(image_data)
        else:
            img = Image.open(io.BytesIO(image_data))

        # Handle transparency for JPEG
        if target_format == ImageFormat.JPEG and img.mode in ("RGBA", "LA", "P"):
            background = Image.new("RGB", img.size, (255, 255, 255))
            if img.mode == "P":
                img = img.convert("RGBA")
            background.paste(img, mask=img.split()[-1])
            img = background

        output = io.BytesIO()
        img.save(output, format=target_format.value.upper())
        return output.getvalue()

    async def extract_text(self, image_data: bytes | Path) -> str:
        """Extract text from image using OCR"""
        if not HAS_PIL or not HAS_TESSERACT:
            return "OCR not available (missing PIL or Tesseract)"

        if isinstance(image_data, Path):
            img = Image.open(image_data)
        else:
            img = Image.open(io.BytesIO(image_data))

        return pytesseract.image_to_string(img)


class DocumentProcessor:
    """Process and extract content from documents"""

    def __init__(self):
        self.supported_formats = list(DocumentFormat)

    async def extract(
        self,
        document_path: str | Path,
        extract_tables: bool = True,
        extract_images: bool = False,
    ) -> DocumentContent:
        """Extract content from a document"""
        path = Path(document_path)
        suffix = path.suffix.lower().lstrip(".")

        if suffix == "pdf":
            return await self._extract_pdf(path, extract_tables, extract_images)
        elif suffix == "docx":
            return await self._extract_docx(path, extract_tables, extract_images)
        elif suffix in ("txt", "md"):
            return await self._extract_text(path)
        elif suffix == "csv":
            return await self._extract_csv(path)
        elif suffix == "json":
            return await self._extract_json(path)
        else:
            # Try as plain text
            return await self._extract_text(path)

    async def _extract_pdf(self, path: Path, extract_tables: bool, extract_images: bool) -> DocumentContent:
        """Extract content from PDF"""
        if not HAS_PYMUPDF:
            return DocumentContent(
                text="PDF extraction not available (PyMuPDF not installed)",
                pages=0,
                word_count=0,
                format="pdf",
            )

        doc = fitz.open(path)

        text_parts = []
        sections = []
        images = []

        for page_num, page in enumerate(doc):
            text = page.get_text()
            text_parts.append(text)

            sections.append(
                {
                    "page": page_num + 1,
                    "text": text[:500] + "..." if len(text) > 500 else text,
                }
            )

            if extract_images:
                for img_index, img in enumerate(page.get_images()):
                    xref = img[0]
                    base_image = doc.extract_image(xref)
                    image_bytes = base_image["image"]
                    images.append(base64.b64encode(image_bytes).decode())

        full_text = "\n\n".join(text_parts)

        return DocumentContent(
            text=full_text,
            pages=len(doc),
            word_count=len(full_text.split()),
            format="pdf",
            metadata={
                "title": doc.metadata.get("title", ""),
                "author": doc.metadata.get("author", ""),
            },
            sections=sections,
            images=images[:10],  # Limit images
        )

    async def _extract_docx(self, path: Path, extract_tables: bool, extract_images: bool) -> DocumentContent:
        """Extract content from DOCX"""
        if not HAS_DOCX:
            return DocumentContent(
                text="DOCX extraction not available (python-docx not installed)",
                pages=0,
                word_count=0,
                format="docx",
            )

        doc = DocxDocument(path)

        text_parts = []
        sections = []
        tables = []

        for para in doc.paragraphs:
            text_parts.append(para.text)
            if para.style.name.startswith("Heading"):
                sections.append({"type": para.style.name, "text": para.text})

        if extract_tables:
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)

        full_text = "\n".join(text_parts)

        return DocumentContent(
            text=full_text,
            pages=1,  # DOCX doesn't have clear page boundaries
            word_count=len(full_text.split()),
            format="docx",
            sections=sections,
            tables=tables,
        )

    async def _extract_text(self, path: Path) -> DocumentContent:
        """Extract content from plain text file"""
        with open(path, encoding="utf-8", errors="replace") as f:
            text = f.read()

        return DocumentContent(
            text=text,
            pages=1,
            word_count=len(text.split()),
            format=path.suffix.lstrip("."),
        )

    async def _extract_csv(self, path: Path) -> DocumentContent:
        """Extract content from CSV"""
        if HAS_PANDAS:
            df = pd.read_csv(path)
            text = df.to_string()
            tables = [df.columns.tolist()] + df.values.tolist()
        else:
            import csv

            with open(path, encoding="utf-8") as f:
                reader = csv.reader(f)
                tables = [list(row) for row in reader]
            text = "\n".join([",".join(row) for row in tables])

        return DocumentContent(
            text=text,
            pages=1,
            word_count=len(text.split()),
            format="csv",
            tables=[tables],
        )

    async def _extract_json(self, path: Path) -> DocumentContent:
        """Extract content from JSON"""
        with open(path, encoding="utf-8") as f:
            data = json.load(f)

        text = json.dumps(data, indent=2)

        return DocumentContent(
            text=text,
            pages=1,
            word_count=len(text.split()),
            format="json",
            metadata={"type": type(data).__name__},
        )


class CodeAnalyzer:
    """Analyze source code"""

    LANGUAGE_EXTENSIONS = {
        ".py": "python",
        ".js": "javascript",
        ".ts": "typescript",
        ".jsx": "javascript",
        ".tsx": "typescript",
        ".java": "java",
        ".cpp": "cpp",
        ".c": "c",
        ".go": "go",
        ".rs": "rust",
        ".rb": "ruby",
        ".php": "php",
        ".swift": "swift",
        ".kt": "kotlin",
    }

    async def analyze(
        self, code: str, language: str | None = None, file_path: str | None = None
    ) -> CodeAnalysisResult:
        """Analyze source code"""
        # Detect language
        if not language and file_path:
            ext = Path(file_path).suffix.lower()
            language = self.LANGUAGE_EXTENSIONS.get(ext, "unknown")
        language = language or "unknown"

        # Count lines
        lines = code.split("\n")
        loc = len([line for line in lines if line.strip() and not line.strip().startswith("#")])

        # Language-specific analysis
        if language == "python":
            return await self._analyze_python(code, loc)
        else:
            return await self._analyze_generic(code, language, loc)

    async def _analyze_python(self, code: str, loc: int) -> CodeAnalysisResult:
        """Analyze Python code"""
        functions = []
        classes = []
        imports = []
        issues = []
        suggestions = []

        if HAS_AST:
            try:
                tree = ast.parse(code)
                for node in ast.walk(tree):
                    if isinstance(node, ast.FunctionDef):
                        functions.append(
                            {
                                "name": node.name,
                                "line": node.lineno,
                                "args": [arg.arg for arg in node.args.args],
                                "has_docstring": ast.get_docstring(node) is not None,
                                "decorators": [
                                    d.id if isinstance(d, ast.Name) else str(d) for d in node.decorator_list
                                ],
                            }
                        )
                    elif isinstance(node, ast.AsyncFunctionDef):
                        functions.append(
                            {
                                "name": node.name,
                                "line": node.lineno,
                                "args": [arg.arg for arg in node.args.args],
                                "has_docstring": ast.get_docstring(node) is not None,
                                "async": True,
                            }
                        )
                    elif isinstance(node, ast.ClassDef):
                        classes.append(
                            {
                                "name": node.name,
                                "line": node.lineno,
                                "bases": [b.id if isinstance(b, ast.Name) else str(b) for b in node.bases],
                                "has_docstring": ast.get_docstring(node) is not None,
                                "methods": len(
                                    [n for n in node.body if isinstance(n, (ast.FunctionDef, ast.AsyncFunctionDef))]
                                ),
                            }
                        )
                    elif isinstance(node, ast.Import):
                        imports.extend([alias.name for alias in node.names])
                    elif isinstance(node, ast.ImportFrom):
                        if node.module:
                            imports.append(node.module)

                # Check for issues
                for node in ast.walk(tree):
                    if isinstance(node, ast.FunctionDef):
                        if not ast.get_docstring(node):
                            issues.append(
                                {
                                    "type": "missing_docstring",
                                    "line": node.lineno,
                                    "message": f"Function '{node.name}' lacks docstring",
                                }
                            )
                        if len(node.args.args) > 5:
                            issues.append(
                                {
                                    "type": "too_many_args",
                                    "line": node.lineno,
                                    "message": f"Function '{node.name}' has too many arguments",
                                }
                            )

            except SyntaxError as e:
                issues.append({"type": "syntax_error", "line": e.lineno, "message": str(e)})

        # Calculate complexity (simplified)
        complexity = 1.0
        complexity += len(functions) * 0.1
        complexity += len(classes) * 0.2
        complexity += code.count("if ") * 0.05
        complexity += code.count("for ") * 0.05
        complexity += code.count("while ") * 0.05
        complexity += code.count("try:") * 0.1

        # Documentation coverage
        documented = len([f for f in functions if f.get("has_docstring")])
        doc_coverage = documented / len(functions) if functions else 1.0

        # Suggestions
        if doc_coverage < 0.5:
            suggestions.append("Consider adding docstrings to more functions")
        if len(imports) > 20:
            suggestions.append("Consider organizing imports or splitting the module")
        if loc > 500:
            suggestions.append("Consider splitting this file into smaller modules")

        return CodeAnalysisResult(
            language="python",
            lines_of_code=loc,
            functions=functions,
            classes=classes,
            imports=imports,
            complexity_score=min(complexity, 10.0),
            issues=issues,
            suggestions=suggestions,
            documentation_coverage=doc_coverage,
        )

    async def _analyze_generic(self, code: str, language: str, loc: int) -> CodeAnalysisResult:
        """Generic code analysis for unsupported languages"""
        # Simple pattern-based analysis
        functions = []
        classes = []
        imports = []

        # Function patterns
        func_patterns = {
            "javascript": r"(?:function\s+(\w+)|(\w+)\s*=\s*(?:async\s+)?function|(\w+)\s*=\s*(?:async\s+)?\([^)]*\)\s*=>)",
            "java": r"(?:public|private|protected)?\s*(?:static)?\s*\w+\s+(\w+)\s*\([^)]*\)",
            "go": r"func\s+(?:\([^)]+\)\s+)?(\w+)\s*\(",
        }

        pattern = func_patterns.get(language)
        if pattern:
            for match in re.finditer(pattern, code):
                name = next((g for g in match.groups() if g), "anonymous")
                functions.append({"name": name, "line": code[: match.start()].count("\n") + 1})

        # Class patterns
        class_patterns = {
            "javascript": r"class\s+(\w+)",
            "java": r"class\s+(\w+)",
            "typescript": r"class\s+(\w+)",
        }

        pattern = class_patterns.get(language)
        if pattern:
            for match in re.finditer(pattern, code):
                classes.append(
                    {
                        "name": match.group(1),
                        "line": code[: match.start()].count("\n") + 1,
                    }
                )

        return CodeAnalysisResult(
            language=language,
            lines_of_code=loc,
            functions=functions,
            classes=classes,
            imports=imports,
            complexity_score=1.0 + len(functions) * 0.1 + len(classes) * 0.2,
            issues=[],
            suggestions=[],
            documentation_coverage=0.0,
        )

    async def generate_documentation(self, code: str, language: str = "python") -> str:
        """Generate documentation for code"""
        analysis = await self.analyze(code, language)

        doc_parts = [
            f"# Code Documentation\n\n**Language:** {analysis.language}\n**Lines of Code:** {analysis.lines_of_code}\n"
        ]

        if analysis.classes:
            doc_parts.append("\n## Classes\n")
            for cls in analysis.classes:
                doc_parts.append(f"- **{cls['name']}** (line {cls.get('line', '?')})")
                if cls.get("bases"):
                    doc_parts.append(f"  - Inherits from: {', '.join(cls['bases'])}")
                if cls.get("methods"):
                    doc_parts.append(f"  - Methods: {cls['methods']}")

        if analysis.functions:
            doc_parts.append("\n## Functions\n")
            for func in analysis.functions:
                doc_parts.append(f"- **{func['name']}** (line {func.get('line', '?')})")
                if func.get("args"):
                    doc_parts.append(f"  - Arguments: {', '.join(func['args'])}")
                if func.get("async"):
                    doc_parts.append("  - Async: Yes")

        if analysis.imports:
            doc_parts.append("\n## Dependencies\n")
            for imp in analysis.imports[:20]:
                doc_parts.append(f"- {imp}")

        if analysis.issues:
            doc_parts.append("\n## Issues Found\n")
            for issue in analysis.issues[:10]:
                doc_parts.append(f"- Line {issue.get('line', '?')}: {issue['message']}")

        return "\n".join(doc_parts)


class DataVisualizer:
    """Create data visualizations"""

    def __init__(self):
        self.default_style = "seaborn-v0_8-whitegrid"

    async def create_chart(
        self,
        data: dict[str, Any],
        chart_type: str = "bar",
        title: str = "",
        x_label: str = "",
        y_label: str = "",
        width: int = 10,
        height: int = 6,
    ) -> bytes:
        """Create a chart from data"""
        if not HAS_MATPLOTLIB:
            raise RuntimeError("Matplotlib not installed")

        plt.figure(figsize=(width, height))

        try:
            plt.style.use(self.default_style)
        except Exception as e:
            logger.debug("Failed to apply matplotlib style %r: %s", self.default_style, e)

        x = data.get("x", data.get("labels", list(range(len(data.get("y", []))))))
        y = data.get("y", data.get("values", []))

        if chart_type == "bar":
            plt.bar(x, y, color=data.get("color", "steelblue"))
        elif chart_type == "line":
            plt.plot(x, y, marker="o", color=data.get("color", "steelblue"))
        elif chart_type == "scatter":
            plt.scatter(x, y, color=data.get("color", "steelblue"))
        elif chart_type == "pie":
            plt.pie(y, labels=x, autopct="%1.1f%%")
        elif chart_type == "histogram":
            plt.hist(y, bins=data.get("bins", 10), color=data.get("color", "steelblue"))
        else:
            plt.bar(x, y)

        if title:
            plt.title(title)
        if x_label:
            plt.xlabel(x_label)
        if y_label:
            plt.ylabel(y_label)

        plt.tight_layout()

        # Save to bytes
        buf = io.BytesIO()
        plt.savefig(buf, format="png", dpi=100)
        plt.close()
        buf.seek(0)

        return buf.getvalue()

    async def create_table_image(
        self,
        data: list[list[str]],
        headers: list[str] | None = None,
        title: str = "",
    ) -> bytes:
        """Create an image of a table"""
        if not HAS_MATPLOTLIB:
            raise RuntimeError("Matplotlib not installed")

        fig, ax = plt.subplots(figsize=(12, len(data) * 0.5 + 1))
        ax.axis("off")

        if headers:
            data = [headers] + data

        table = ax.table(cellText=data, loc="center", cellLoc="center")

        table.auto_set_font_size(False)
        table.set_fontsize(10)
        table.scale(1.2, 1.5)

        if title:
            plt.title(title, fontsize=14, fontweight="bold")

        buf = io.BytesIO()
        plt.savefig(buf, format="png", dpi=100, bbox_inches="tight")
        plt.close()
        buf.seek(0)

        return buf.getvalue()

    async def dataframe_to_chart(
        self,
        df_json: str,
        x_column: str,
        y_column: str,
        chart_type: str = "bar",
        title: str = "",
    ) -> bytes:
        """Create chart from DataFrame JSON"""
        if not HAS_PANDAS:
            raise RuntimeError("Pandas not installed")

        df = pd.read_json(df_json)

        data = {"x": df[x_column].tolist(), "y": df[y_column].tolist()}

        return await self.create_chart(data, chart_type, title, x_column, y_column)


class MultimodalProcessor:
    """
    Main multimodal processor that combines all capabilities.
    Provides a unified interface for image, document, code, and data processing.
    """

    def __init__(self):
        self.image_processor = ImageProcessor()
        self.document_processor = DocumentProcessor()
        self.code_analyzer = CodeAnalyzer()
        self.data_visualizer = DataVisualizer()

    # Image processing methods
    async def analyze_image(self, image_path: str | Path) -> dict[str, Any]:
        """Analyze an image"""
        return await self.image_processor.analyze(image_path)

    async def resize_image(
        self,
        image_data: bytes | Path,
        width: int | None = None,
        height: int | None = None,
    ) -> bytes:
        """Resize an image"""
        return await self.image_processor.resize(image_data, width, height)

    async def convert_image(self, image_data: bytes | Path, target_format: ImageFormat) -> bytes:
        """Convert image format"""
        return await self.image_processor.convert(image_data, target_format)

    async def extract_text_from_image(self, image_data: bytes | Path) -> str:
        """Extract text from image using OCR"""
        return await self.image_processor.extract_text(image_data)

    # Document processing methods
    async def process_document(
        self,
        document_path: str | Path,
        extract_tables: bool = True,
        extract_images: bool = False,
    ) -> DocumentContent:
        """Process a document and extract content"""
        return await self.document_processor.extract(document_path, extract_tables, extract_images)

    # Code analysis methods
    async def analyze_code(
        self,
        code: str,
        language: str | None = None,
        file_path: str | None = None,
    ) -> CodeAnalysisResult:
        """Analyze source code"""
        return await self.code_analyzer.analyze(code, language, file_path)

    async def generate_code_documentation(self, code: str, language: str = "python") -> str:
        """Generate documentation for code"""
        return await self.code_analyzer.generate_documentation(code, language)

    # Data visualization methods
    async def create_chart(
        self,
        data: dict[str, Any],
        chart_type: str = "bar",
        title: str = "",
        x_label: str = "",
        y_label: str = "",
    ) -> bytes:
        """Create a data visualization chart"""
        return await self.data_visualizer.create_chart(data, chart_type, title, x_label, y_label)

    async def create_table_image(
        self,
        data: list[list[str]],
        headers: list[str] | None = None,
        title: str = "",
    ) -> bytes:
        """Create an image of a data table"""
        return await self.data_visualizer.create_table_image(data, headers, title)

    async def create_dataframe_chart(
        self,
        df_json: str,
        x_column: str,
        y_column: str,
        chart_type: str = "bar",
        title: str = "",
    ) -> bytes:
        """Create chart from DataFrame JSON"""
        return await self.data_visualizer.dataframe_to_chart(df_json, x_column, y_column, chart_type, title)

    # Additional methods for testing
    async def generate_visualization(self, data: dict[str, Any], viz_type: str = "chart", **kwargs) -> bytes:
        """Generate data visualization"""
        if viz_type == "chart":
            return await self.create_chart(data, **kwargs)
        elif viz_type == "table":
            return await self.create_table_image(data.get("table_data", []), data.get("headers"), data.get("title", ""))
        else:
            raise ValueError(f"Unsupported visualization type: {viz_type}")

    async def transcribe_audio(self, audio_path: str | Path) -> str:
        """
        Transcribe audio to text using OpenAI Whisper API.

        Supports: mp3, mp4, mpeg, mpga, m4a, wav, webm
        Cost: ~$0.006 per minute of audio
        """
        import openai
        from openai import OpenAI

        # Get API key from environment
        api_key = os.getenv("OPENAI_API_KEY")
        if not api_key:
            logger.warning("OPENAI_API_KEY not set, audio transcription unavailable")
            if not FEATURE_AUDIO_TRANSCRIPTION:
                return "[Audio Transcription] API key not configured. Set OPENAI_API_KEY environment variable."
            raise ValueError("OPENAI_API_KEY environment variable not set")

        try:
            # Initialize OpenAI client
            client = OpenAI(api_key=api_key)

            # Open audio file
            audio_file_path = Path(audio_path)
            if not audio_file_path.exists():
                raise FileNotFoundError(f"Audio file not found: {audio_path}")

            # Check file size (OpenAI limit: 25MB)
            file_size_mb = audio_file_path.stat().st_size / (1024 * 1024)
            if file_size_mb > 25:
                logger.warning("Audio file %s is %.1fMB (limit: 25MB)", audio_path, file_size_mb)
                raise ValueError(
                    f"Audio file too large: {file_size_mb:.1f}MB (max: 25MB). Consider splitting the file."
                )

            # Transcribe using Whisper
            with open(audio_file_path, "rb") as audio_file:
                start_time = datetime.now(UTC)
                transcript = client.audio.transcriptions.create(
                    model="whisper-1", file=audio_file, response_format="text"
                )
                duration = (datetime.now(UTC) - start_time).total_seconds()

            # Log usage for cost tracking
            logger.info(
                f"🌀 Audio transcribed: {audio_file_path.name} "
                f"({file_size_mb:.2f}MB, {duration:.1f}s processing) "
                f"→ {len(transcript)} characters"
            )

            return transcript

        except openai.OpenAIError as e:
            logger.error("OpenAI API error during transcription: %s", e)
            raise RuntimeError(f"Audio transcription failed: {e!s}")
        except Exception as e:
            logger.error("Unexpected error during audio transcription: %s", e)
            raise

    async def generate_image(
        self,
        prompt: str,
        width: int = 1024,
        height: int = 1024,
        style: str = "vivid",
        provider: str = "auto",
    ) -> bytes:
        """
        Generate image from text prompt.

        Providers: "dalle" (DALL-E 3), "grok" (Grok), "auto" (prefers Grok, falls back to DALL-E).

        Args:
            prompt: Text description of image to generate
            width: Image width (1024, 1792 for landscape)
            height: Image height (1024, 1792 for portrait)
            style: "vivid" (hyper-real) or "natural" (more balanced) — DALL-E only
            provider: "auto", "dalle", or "grok"

        Supported sizes (DALL-E): 1024x1024, 1792x1024, 1024x1792
        Grok uses aspect_ratio instead of pixel sizes.
        """
        # Decide which provider to use
        has_xai = bool(os.getenv("XAI_API_KEY"))
        has_openai = bool(os.getenv("OPENAI_API_KEY"))

        if provider == "auto":
            provider = "grok" if has_xai else "dalle"
        if provider == "grok" and not has_xai:
            provider = "dalle"
        if provider == "dalle" and not has_openai:
            if has_xai:
                provider = "grok"

        # --- Grok path ---
        if provider == "grok":
            # Map pixel dimensions to aspect ratio
            if width > height:
                aspect = "16:9"
            elif height > width:
                aspect = "9:16"
            else:
                aspect = "1:1"
            images = await self.generate_image_grok(
                prompt=prompt,
                n=1,
                aspect_ratio=aspect,
            )
            if images:
                return images[0]
            raise RuntimeError("Grok image generation returned no images")

        # --- DALL-E path (original) ---
        import httpx
        import openai
        from openai import OpenAI

        # Get API key from environment
        api_key = os.getenv("OPENAI_API_KEY")
        if not api_key:
            logger.warning("OPENAI_API_KEY not set, image generation unavailable")
            if not FEATURE_IMAGE_GENERATION:
                # Return placeholder image
                if not HAS_PIL:
                    raise RuntimeError("PIL not available and OpenAI API key not configured")
                img = Image.new("RGB", (width, height), color=(100, 150, 200))
                output = io.BytesIO()
                img.save(output, format="PNG")
                return output.getvalue()
            raise ValueError("OPENAI_API_KEY environment variable not set")

        try:
            # Initialize OpenAI client
            client = OpenAI(api_key=api_key)

            # Validate dimensions and map to DALL-E 3 sizes
            size_map = {
                (1024, 1024): "1024x1024",
                (1792, 1024): "1792x1024",
                (1024, 1792): "1024x1792",
            }
            size_key = (width, height)

            if size_key not in size_map:
                logger.warning("Unsupported size %sx%s, using 1024x1024", width, height)
                size_str = "1024x1024"
            else:
                size_str = size_map[size_key]

            # Validate style
            if style not in ["vivid", "natural"]:
                logger.warning("Unsupported style '%s', using 'vivid'", style)
                style = "vivid"

            # Generate image
            start_time = datetime.now(UTC)
            response = client.images.generate(
                model="dall-e-3",
                prompt=prompt,
                size=size_str,
                style=style,
                quality="standard",  # "standard" or "hd" (2x cost)
                n=1,
            )
            duration = (datetime.now(UTC) - start_time).total_seconds()

            # Download the image
            image_url = response.data[0].url
            async with httpx.AsyncClient(timeout=30.0) as http_client:
                image_response = await http_client.get(image_url)
                image_response.raise_for_status()
                image_bytes = image_response.content

            # Log usage for cost tracking
            cost_estimate = 0.040 if size_str == "1024x1024" else 0.080
            logger.info(
                f"🌀 Image generated: '{prompt[:50]}...' "
                f"({size_str}, {style}, {duration:.1f}s processing) "
                f"→ {len(image_bytes)/1024:.1f}KB (est. ${cost_estimate:.3f})"
            )

            return image_bytes

        except openai.OpenAIError as e:
            logger.error("OpenAI API error during image generation: %s", e)
            raise RuntimeError(f"Image generation failed: {e!s}")
        except Exception as e:
            logger.error("Unexpected error during image generation: %s", e)
            raise

    async def generate_image_grok(
        self,
        prompt: str,
        n: int = 1,
        aspect_ratio: str = "1:1",
        response_format: str = "b64_json",
    ) -> list[bytes]:
        """
        Generate images using xAI Grok (grok-imagine-image).

        Args:
            prompt: Text description of image to generate
            n: Number of images (1-10)
            aspect_ratio: e.g. "1:1", "16:9", "9:16", "4:3"
            response_format: "b64_json" or "url"

        Returns list of image byte arrays (PNG).
        """
        import httpx

        api_key = os.getenv("XAI_API_KEY")
        if not api_key:
            raise ValueError("XAI_API_KEY environment variable not set")

        n = max(1, min(n, 10))
        valid_ratios = [
            "1:1",
            "16:9",
            "9:16",
            "4:3",
            "3:4",
            "3:2",
            "2:3",
            "2:1",
            "1:2",
            "auto",
        ]
        if aspect_ratio not in valid_ratios:
            aspect_ratio = "1:1"

        start_time = datetime.now(UTC)

        async with httpx.AsyncClient(timeout=120.0) as client:
            response = await client.post(
                "https://api.x.ai/v1/images/generations",
                headers={
                    "Authorization": f"Bearer {api_key}",
                    "Content-Type": "application/json",
                },
                json={
                    "model": "grok-imagine-image",
                    "prompt": prompt,
                    "n": n,
                    "response_format": response_format,
                },
            )
            response.raise_for_status()
            data = response.json()

        duration = (datetime.now(UTC) - start_time).total_seconds()
        images: list[bytes] = []

        for item in data.get("data", []):
            if response_format == "b64_json" and item.get("b64_json"):
                images.append(base64.b64decode(item["b64_json"]))
            elif item.get("url"):
                async with httpx.AsyncClient(timeout=60.0) as dl:
                    img_resp = await dl.get(item["url"])
                    img_resp.raise_for_status()
                    images.append(img_resp.content)

        logger.info(
            "Grok image generated: '%s' (%s, n=%d, %.1fs) -> %d images",
            prompt[:50],
            aspect_ratio,
            n,
            duration,
            len(images),
        )
        return images

    async def generate_video_grok(
        self,
        prompt: str,
        duration: int = 5,
        aspect_ratio: str = "16:9",
        resolution: str = "720p",
        image_url: str | None = None,
    ) -> dict[str, Any]:
        """
        Generate a video using xAI Grok (grok-imagine-video).

        Args:
            prompt: Text description of the video
            duration: 1-15 seconds
            aspect_ratio: "1:1", "16:9", "9:16", etc.
            resolution: "480p" or "720p"
            image_url: Optional source image for image-to-video

        Returns dict with 'video_url', 'duration', and metadata.
        """
        import httpx

        api_key = os.getenv("XAI_API_KEY")
        if not api_key:
            raise ValueError("XAI_API_KEY environment variable not set")

        duration = max(1, min(duration, 15))
        if resolution not in ("480p", "720p"):
            resolution = "720p"

        body: dict[str, Any] = {
            "model": "grok-imagine-video",
            "prompt": prompt,
            "duration": duration,
            "aspect_ratio": aspect_ratio,
            "resolution": resolution,
        }
        if image_url:
            body["image_url"] = image_url

        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
        }

        # Submit generation request
        async with httpx.AsyncClient(timeout=30.0) as client:
            submit_resp = await client.post(
                "https://api.x.ai/v1/videos/generations",
                headers=headers,
                json=body,
            )
            submit_resp.raise_for_status()
            submit_data = submit_resp.json()

        request_id = submit_data.get("request_id")
        if not request_id:
            raise RuntimeError("No request_id returned from video generation API")

        # Poll for completion (max ~5 minutes)
        import asyncio

        poll_url = f"https://api.x.ai/v1/videos/{request_id}"
        for _ in range(300):  # 300 * 1s = 5 min max
            await asyncio.sleep(1.0)
            async with httpx.AsyncClient(timeout=30.0) as client:
                poll_resp = await client.get(poll_url, headers=headers)
                poll_resp.raise_for_status()
                poll_data = poll_resp.json()

            status = poll_data.get("status")
            if status == "done":
                video_info = poll_data.get("video", {})
                logger.info(
                    "Grok video generated: '%s' (%ds, %s, %s)",
                    prompt[:50],
                    duration,
                    aspect_ratio,
                    resolution,
                )
                return {
                    "video_url": video_info.get("url"),
                    "duration": video_info.get("duration"),
                    "status": "done",
                    "request_id": request_id,
                }
            elif status == "expired":
                raise RuntimeError("Video generation request expired")
            # else: still pending, continue polling

        raise RuntimeError("Video generation timed out after 5 minutes")

    def get_supported_formats(self) -> list[str]:
        """Get supported formats for different operations"""
        all_formats = []
        all_formats.extend(["png", "jpeg", "gif", "webp", "bmp"])  # images
        all_formats.extend(["pdf", "docx", "txt", "md", "html", "csv", "json", "xml"])  # documents
        all_formats.extend(["wav", "mp3", "flac", "ogg"])  # audio
        all_formats.extend(["mp4", "avi", "mov"])  # video
        return all_formats
